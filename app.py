# app.py
# Aplikasi Register Pasien IGD - UPDATED VERSION

from flask import Flask, render_template_string, request, redirect, session, send_file
from flask_sqlalchemy import SQLAlchemy
from docx import Document
from datetime import datetime
import os

app = Flask(__name__)
app.secret_key = "secretkey123"
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv("DATABASE_URL")
if app.config['SQLALCHEMY_DATABASE_URI'].startswith("postgres://"):
    app.config['SQLALCHEMY_DATABASE_URI'] = app.config['SQLALCHEMY_DATABASE_URI'].replace(
        "postgres://", "postgresql://", 1
    )
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# ================= DATABASE =================
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True)
    password = db.Column(db.String(100))

class Patient(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    nama = db.Column(db.String(200))
    umur = db.Column(db.String(50))
    jenis_kelamin = db.Column(db.String(50))
    alamat = db.Column(db.String(300))
    soap_s = db.Column(db.Text)
    soap_o = db.Column(db.Text)
    soap_a = db.Column(db.Text)
    soap_p = db.Column(db.Text)
    status_lokalis = db.Column(db.Text)
    created_by = db.Column(db.String(100))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

# ================= HOME =================
@app.route('/')
def home():
    if 'user' in session:
        return redirect('/register_patient')
    return redirect('/login')

# ================= REGISTER USER =================
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        user = User(username=request.form['username'], password=request.form['password'])
        db.session.add(user)
        db.session.commit()
        return redirect('/login')

    return render_template_string('''
    <h2>Register User</h2>
    <form method="post">
        Username: <input name="username"><br><br>
        Password: <input type="password" name="password"><br><br>
        <button type="submit">Daftar</button>
    </form>
    <a href="/login">Login</a>
    ''')

# ================= LOGIN =================
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        user = User.query.filter_by(
            username=request.form['username'],
            password=request.form['password']
        ).first()
        if user:
            session['user'] = user.username
            return redirect('/register_patient')

    return render_template_string('''
<!doctype html>
<html>
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
<style>
body{
    background: url('https://images.unsplash.com/photo-1588776814546-1ffcf47267a5') no-repeat center center fixed;
    background-size: cover;
}
.overlay{
    background: rgba(255,255,255,0.95);
    border-radius:20px;
}
.footer-text{
    position: fixed;
    bottom: 10px;
    right: 15px;
    font-size: 12px;
    color: #555;
}
</style>
</head>
<body class="d-flex justify-content-center align-items-center vh-100">
<div class="container">
<div class="row justify-content-center">
<div class="col-md-5">
<div class="overlay shadow-lg p-5">
<div class="text-center mb-4">
<h2 class="fw-bold text-danger">🩺 PUSKESMAS BATULICIN</h2>
<p class="text-muted">Sistem Register Pasien IGD</p>
<hr>
</div>
<form method="post">
<input name="username" class="form-control mb-3" placeholder="Username" required>
<input type="password" name="password" class="form-control mb-3" placeholder="Password" required>
<button class="btn btn-success w-100">Login</button>
</form>
<div class="text-center mt-3">
<a href="/register">Daftar User</a>
</div>
</div>
</div>
</div>
</div>

<div class="footer-text">
dibuat oleh dr.M Iqbal Al Islamy 2026
</div>

</body>
</html>
''')

# ================= LOGOUT =================
@app.route('/logout')
def logout():
    session.pop('user', None)
    return redirect('/login')

# ================= REGISTER PATIENT =================
@app.route('/register_patient', methods=['GET', 'POST'])
def register_patient():
    if 'user' not in session:
        return redirect('/login')

    if request.method == 'POST':

        soap_s = f"""Keluhan Utama: {request.form['keluhan_utama']}
RPS: {request.form['rps']}
RPD: {request.form['rpd']}
RMO: {request.form['rmo']}"""

        soap_o = f"""GCS: {request.form['gcs']}
TD: {request.form['td']}
HR: {request.form['hr']}
RR: {request.form['rr']}
T: {request.form['t']}
SPO2: {request.form['spo2']}
VAS: {request.form['vas']}
BB: {request.form['bb']}
TB: {request.form['tb']}

Head to Toe:
Kepala/Leher: {request.form['kepala']}
Thorax: {request.form['thorax']}
Abdomen: {request.form['abdomen']}
Ekstremitas: {request.form['ekstremitas']}"""

        patient = Patient(
            nama=request.form['nama'],
            umur=request.form['umur'],
            jenis_kelamin=request.form['jenis_kelamin'],
            alamat=request.form['alamat'],
            soap_s=soap_s,
            soap_o=soap_o,
            soap_a=request.form['soap_a'],
            soap_p=request.form['soap_p'],
            status_lokalis=request.form['status_lokalis'],
            created_by=session['user']
        )

        db.session.add(patient)
        db.session.commit()
        return redirect(f"/export/{patient.id}")

    return render_template_string('''
<!doctype html>
<html>
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
<style>
body{
    background: url('https://images.unsplash.com/photo-1579684385127-1ef15d508118') no-repeat center center fixed;
    background-size: cover;
}
.card-custom{
    background: rgba(255,255,255,0.95);
    border-radius:25px;
}
.footer-text{
    position: fixed;
    bottom: 10px;
    right: 15px;
    font-size: 12px;
    color: #555;
}
</style>
</head>
<body>
<div class="container mt-5 mb-5">
<div class="card card-custom shadow-lg p-5">

<div class="text-center mb-4">
<h2 class="fw-bold text-danger">🩺 PUSKESMAS BATULICIN</h2>
<p class="text-secondary fs-5">Sistem Register Pasien IGD</p>
<hr>
</div>

<form method="post">

<h5 class="text-primary">Identitas Pasien</h5>
<input name="nama" class="form-control mb-2" placeholder="Nama" required>
<input name="umur" class="form-control mb-2" placeholder="Umur" required>
<select name="jenis_kelamin" class="form-select mb-2">
<option>Laki-laki</option>
<option>Perempuan</option>
</select>
<input name="alamat" class="form-control mb-3" placeholder="Alamat">

<h5 class="text-primary">S (Subjective)</h5>
<input name="keluhan_utama" class="form-control mb-2" placeholder="Keluhan Utama">
<textarea name="rps" class="form-control mb-2" placeholder="Riwayat Penyakit Sekarang"></textarea>
<textarea name="rpd" class="form-control mb-2" placeholder="Riwayat Penyakit Dahulu"></textarea>
<textarea name="rmo" class="form-control mb-3" placeholder="Riwayat Minum Obat"></textarea>

<h5 class="text-primary">O (Objective)</h5>
<input name="gcs" class="form-control mb-2" placeholder="GCS">
<input name="td" class="form-control mb-2" placeholder="TD">
<input name="hr" class="form-control mb-2" placeholder="HR">
<input name="rr" class="form-control mb-2" placeholder="RR">
<input name="t" class="form-control mb-2" placeholder="T">
<input name="spo2" class="form-control mb-2" placeholder="SPO2">
<input name="vas" class="form-control mb-2" placeholder="VAS">
<input name="bb" class="form-control mb-2" placeholder="BB">
<input name="tb" class="form-control mb-2" placeholder="TB">

<input name="kepala" class="form-control mb-2" placeholder="Kepala/Leher (+/-)">
<input name="thorax" class="form-control mb-2" placeholder="Thorax (+/-)">
<input name="abdomen" class="form-control mb-2" placeholder="Abdomen (+/-)">
<input name="ekstremitas" class="form-control mb-3" placeholder="Ekstremitas (+/-)">

<h5 class="text-primary">Status Lokalis</h5>
<textarea name="status_lokalis" class="form-control mb-3"></textarea>

<h5 class="text-primary">A (Assessment)</h5>
<textarea name="soap_a" class="form-control mb-3"></textarea>

<h5 class="text-primary">P (Plan)</h5>
<textarea name="soap_p" class="form-control mb-4"></textarea>

<div class="d-flex gap-2">
<button class="btn btn-success">💾 Simpan & Export</button>
<button type="reset" class="btn btn-warning">🔄 Reset</button>
<a href="/logout" class="btn btn-secondary ms-auto">Logout</a>
</div>

</form>
</div>
</div>

<div class="footer-text">
dibuat oleh dr.M Iqbal Al Islamy 2026
</div>

</body>
</html>
''')

# ================= EXPORT WORD =================
@app.route('/export/<int:id>')
def export(id):
    patient = Patient.query.get_or_404(id)

    document = Document()
    document.add_heading('REGISTER PASIEN IGD', 0)

    document.add_paragraph(f"Nama: {patient.nama}")
    document.add_paragraph(f"Umur: {patient.umur}")
    document.add_paragraph(f"Jenis Kelamin: {patient.jenis_kelamin}")
    document.add_paragraph(f"Alamat: {patient.alamat}")

    document.add_heading('S (Subjective)', level=1)
    document.add_paragraph(patient.soap_s or "")

    document.add_heading('O (Objective)', level=1)
    document.add_paragraph(patient.soap_o or "")

    document.add_heading('Status Lokalis', level=1)
    document.add_paragraph(patient.status_lokalis or "")

    document.add_heading('A (Assessment)', level=1)
    document.add_paragraph(patient.soap_a or "")

    document.add_heading('P (Plan)', level=1)
    document.add_paragraph(patient.soap_p or "")

    document.add_paragraph("")
    document.add_paragraph(f"Diisi oleh: {patient.created_by}")
    document.add_paragraph(f"Tanggal: {patient.created_at.strftime('%d-%m-%Y %H:%M')}")

    safe_name = patient.nama.replace(" ", "_")
    filename = f"{safe_name}.docx"
    document.save(filename)

    return send_file(filename, as_attachment=True)

# ===== Railway / Gunicorn entry =====
import os

if __name__ == "__main__":
    with app.app_context():
        db.create_all()

    port = int(os.environ.get("PORT", 8080))
    app.run(host="0.0.0.0", port=port)